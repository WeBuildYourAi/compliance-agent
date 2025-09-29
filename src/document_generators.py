"""
Document Generation Module for Compliance Agent
Handles creation of DOCX, PDF, and HTML documents
"""
import os
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional

# Try to import document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX generation disabled.")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab not installed. PDF generation disabled.")

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Generate compliance documents in various formats"""
    
    def __init__(self):
        # Ensure output directory exists
        self.output_dir = Path("/mnt/user-data/outputs")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Document generator initialized. Output directory: {self.output_dir}")
    
    def generate_document_package(
        self,
        document_content: Dict[str, Any],
        document_type: str,
        framework: str,
        company_name: str,
        formats: List[str] = ["docx", "pdf"]
    ) -> Dict[str, str]:
        """Generate document package in multiple formats"""
        
        generated_files = {}
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{framework}_{document_type}_{company_name.replace(' ', '_')}_{timestamp}"
        
        # Generate each requested format
        if "docx" in formats and DOCX_AVAILABLE:
            try:
                docx_path = self.generate_docx(document_content, base_filename)
                generated_files["docx"] = str(docx_path)
                logger.info(f"Generated DOCX: {docx_path}")
            except Exception as e:
                logger.error(f"DOCX generation failed: {e}")
        
        if "pdf" in formats and PDF_AVAILABLE:
            try:
                pdf_path = self.generate_pdf(document_content, base_filename)
                generated_files["pdf"] = str(pdf_path)
                logger.info(f"Generated PDF: {pdf_path}")
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")
        
        if "html" in formats:
            try:
                html_path = self.generate_html(document_content, base_filename)
                generated_files["html"] = str(html_path)
                logger.info(f"Generated HTML: {html_path}")
            except Exception as e:
                logger.error(f"HTML generation failed: {e}")
        
        # Always generate JSON as fallback
        if not generated_files:
            json_path = self.generate_json_fallback(document_content, base_filename)
            generated_files["json"] = str(json_path)
            logger.info(f"Generated JSON fallback: {json_path}")
        
        return generated_files
    
    def generate_docx(self, content: Dict[str, Any], base_filename: str) -> Path:
        """Generate DOCX document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        doc = Document()
        
        # Add document metadata
        metadata = content.get("document_metadata", {})
        doc_content = content.get("document_content", {})
        
        # Add title
        title = metadata.get("title", "Compliance Document")
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        doc.add_paragraph(f"Document Type: {metadata.get('document_type', 'N/A')}")
        doc.add_paragraph(f"Version: {metadata.get('version', '1.0')}")
        doc.add_paragraph(f"Created: {metadata.get('created_date', datetime.utcnow().isoformat())}")
        doc.add_paragraph(f"Frameworks: {', '.join(metadata.get('applicable_frameworks', []))}")
        
        doc.add_page_break()
        
        # Add executive summary if present
        if "executive_summary" in doc_content:
            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(doc_content["executive_summary"])
            doc.add_paragraph()
        
        # Add main sections
        for section in doc_content.get("main_sections", []):
            # Add section heading
            doc.add_heading(section.get("section_title", "Section"), 1)
            
            # Add section content
            section_text = section.get("section_content", "")
            if section_text:
                doc.add_paragraph(section_text)
            
            # Add subsections
            for subsection in section.get("subsections", []):
                doc.add_heading(subsection.get("subsection_title", "Subsection"), 2)
                doc.add_paragraph(subsection.get("subsection_content", ""))
            
            doc.add_paragraph()  # Add spacing
        
        # Add appendices if present
        if "appendices" in doc_content and doc_content["appendices"]:
            doc.add_page_break()
            doc.add_heading("Appendices", 1)
            for appendix in doc_content["appendices"]:
                doc.add_heading(appendix.get("appendix_title", "Appendix"), 2)
                doc.add_paragraph(appendix.get("appendix_content", ""))
        
        # Add footer if present
        if "footer_content" in doc_content:
            doc.add_page_break()
            footer_para = doc.add_paragraph(doc_content["footer_content"])
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        output_path = self.output_dir / f"{base_filename}.docx"
        doc.save(str(output_path))
        
        return output_path
    
    def generate_pdf(self, content: Dict[str, Any], base_filename: str) -> Path:
        """Generate PDF document"""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab not installed")
        
        output_path = self.output_dir / f"{base_filename}.pdf"
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#2c5282'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Add content
        metadata = content.get("document_metadata", {})
        doc_content = content.get("document_content", {})
        
        # Add title
        title = metadata.get("title", "Compliance Document")
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 12))
        
        # Add metadata
        elements.append(Paragraph(f"<b>Document Type:</b> {metadata.get('document_type', 'N/A')}", styles['Normal']))
        elements.append(Paragraph(f"<b>Version:</b> {metadata.get('version', '1.0')}", styles['Normal']))
        elements.append(Paragraph(f"<b>Created:</b> {metadata.get('created_date', datetime.utcnow().isoformat())}", styles['Normal']))
        elements.append(Paragraph(f"<b>Frameworks:</b> {', '.join(metadata.get('applicable_frameworks', []))}", styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Add executive summary
        if "executive_summary" in doc_content:
            elements.append(Paragraph("Executive Summary", heading_style))
            elements.append(Paragraph(doc_content["executive_summary"], styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # Add main sections
        for section in doc_content.get("main_sections", []):
            elements.append(Paragraph(section.get("section_title", "Section"), heading_style))
            
            section_text = section.get("section_content", "")
            if section_text:
                # Handle long text by splitting into paragraphs
                for para in section_text.split('\n'):
                    if para.strip():
                        elements.append(Paragraph(para, styles['Normal']))
                        elements.append(Spacer(1, 6))
            
            # Add subsections
            for subsection in section.get("subsections", []):
                elements.append(Paragraph(subsection.get("subsection_title", "Subsection"), styles['Heading2']))
                subsection_text = subsection.get("subsection_content", "")
                if subsection_text:
                    elements.append(Paragraph(subsection_text, styles['Normal']))
                    elements.append(Spacer(1, 6))
        
        # Add appendices
        if "appendices" in doc_content and doc_content["appendices"]:
            elements.append(PageBreak())
            elements.append(Paragraph("Appendices", heading_style))
            for appendix in doc_content["appendices"]:
                elements.append(Paragraph(appendix.get("appendix_title", "Appendix"), styles['Heading2']))
                elements.append(Paragraph(appendix.get("appendix_content", ""), styles['Normal']))
                elements.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(elements)
        
        return output_path
    
    def generate_html(self, content: Dict[str, Any], base_filename: str) -> Path:
        """Generate HTML document"""
        
        metadata = content.get("document_metadata", {})
        doc_content = content.get("document_content", {})
        
        html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('title', 'Compliance Document')}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .document-container {{
            background-color: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #1f4788;
            border-bottom: 3px solid #1f4788;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #2c5282;
            margin-top: 30px;
        }}
        h3 {{
            color: #4a5568;
        }}
        .metadata {{
            background-color: #e6f2ff;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
        .metadata p {{
            margin: 5px 0;
        }}
        .executive-summary {{
            background-color: #f0f9ff;
            padding: 20px;
            border-left: 4px solid #3182ce;
            margin: 20px 0;
        }}
        .section {{
            margin-bottom: 30px;
        }}
        .subsection {{
            margin-left: 20px;
            margin-top: 15px;
        }}
        .appendix {{
            background-color: #f7fafc;
            padding: 15px;
            margin-top: 20px;
            border-radius: 5px;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e2e8f0;
            color: #718096;
        }}
    </style>
</head>
<body>
    <div class="document-container">
        <h1>{metadata.get('title', 'Compliance Document')}</h1>
        
        <div class="metadata">
            <p><strong>Document Type:</strong> {metadata.get('document_type', 'N/A')}</p>
            <p><strong>Version:</strong> {metadata.get('version', '1.0')}</p>
            <p><strong>Created:</strong> {metadata.get('created_date', datetime.utcnow().isoformat())}</p>
            <p><strong>Frameworks:</strong> {', '.join(metadata.get('applicable_frameworks', []))}</p>
        </div>
"""
        
        # Add executive summary
        if "executive_summary" in doc_content:
            html_template += f"""
        <div class="executive-summary">
            <h2>Executive Summary</h2>
            <p>{doc_content['executive_summary']}</p>
        </div>
"""
        
        # Add main sections
        for section in doc_content.get("main_sections", []):
            html_template += f"""
        <div class="section">
            <h2>{section.get('section_title', 'Section')}</h2>
            <p>{section.get('section_content', '')}</p>
"""
            
            # Add subsections
            for subsection in section.get("subsections", []):
                html_template += f"""
            <div class="subsection">
                <h3>{subsection.get('subsection_title', 'Subsection')}</h3>
                <p>{subsection.get('subsection_content', '')}</p>
            </div>
"""
            
            html_template += "</div>"
        
        # Add appendices
        if "appendices" in doc_content and doc_content["appendices"]:
            html_template += "<h2>Appendices</h2>"
            for appendix in doc_content["appendices"]:
                html_template += f"""
        <div class="appendix">
            <h3>{appendix.get('appendix_title', 'Appendix')}</h3>
            <p>{appendix.get('appendix_content', '')}</p>
        </div>
"""
        
        # Add footer
        if "footer_content" in doc_content:
            html_template += f"""
        <div class="footer">
            <p>{doc_content['footer_content']}</p>
        </div>
"""
        
        html_template += """
    </div>
</body>
</html>"""
        
        # Save HTML file
        output_path = self.output_dir / f"{base_filename}.html"
        output_path.write_text(html_template, encoding='utf-8')
        
        return output_path
    
    def generate_json_fallback(self, content: Dict[str, Any], base_filename: str) -> Path:
        """Generate JSON document as fallback"""
        
        output_path = self.output_dir / f"{base_filename}.json"
        
        # Write formatted JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, indent=2, ensure_ascii=False, default=str)
        
        return output_path

# Global instance
document_generator = DocumentGenerator()
