"""
Document Service for Compliance Agent
Handles creation of Word and PDF documents from compliance content
Based on the working implementation from deep-research-agent
"""
import io
import asyncio
import base64
from typing import Dict, Any, List, Optional
from datetime import datetime
import logging
from pathlib import Path

# DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX generation disabled.")

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("ReportLab not installed. PDF generation disabled.")

# Excel generation
try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("openpyxl not installed. Excel generation disabled.")

logger = logging.getLogger(__name__)

class ComplianceDocumentService:
    """Service for creating compliance documents in various formats"""
    
    def __init__(self):
        """Initialize the document service"""
        # Ensure output directory exists
        self.output_dir = Path("/mnt/user-data/outputs")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Document service initialized. Output directory: {self.output_dir}")
    
    async def generate_document_package(
        self,
        document_content: Dict[str, Any],
        document_type: str,
        framework: str,
        company_name: str,
        formats: List[str] = ["docx", "pdf"]
    ) -> Dict[str, str]:
        """Generate document package in multiple formats"""
        
        generated_files = {}
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{framework}_{document_type}_{company_name.replace(' ', '_')}_{timestamp}"
        
        # Generate each requested format
        if "docx" in formats and DOCX_AVAILABLE:
            try:
                docx_path = await self.create_docx_document(document_content, base_filename)
                generated_files["docx"] = str(docx_path)
                logger.info(f"Generated DOCX: {docx_path}")
            except Exception as e:
                logger.error(f"DOCX generation failed: {e}")
        
        if "pdf" in formats and PDF_SUPPORT:
            try:
                pdf_path = await self.create_pdf_document(document_content, base_filename)
                generated_files["pdf"] = str(pdf_path)
                logger.info(f"Generated PDF: {pdf_path}")
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")
        
        if "xlsx" in formats and EXCEL_AVAILABLE:
            try:
                xlsx_path = await self.create_excel_document(document_content, document_type, framework, company_name, base_filename)
                generated_files["xlsx"] = str(xlsx_path)
                logger.info(f"Generated Excel: {xlsx_path}")
            except Exception as e:
                logger.error(f"Excel generation failed: {e}")
        
        # Always generate JSON as fallback
        if not generated_files:
            json_path = await self.create_json_fallback(document_content, base_filename)
            generated_files["json"] = str(json_path)
            logger.info(f"Generated JSON fallback: {json_path}")
        
        return generated_files
    
    async def create_docx_document(
        self,
        content: Dict[str, Any],
        base_filename: str
    ) -> Path:
        """Create a Word document from content data asynchronously"""
        
        # Run in thread pool to avoid blocking
        path = await asyncio.to_thread(
            self._create_docx_document_sync,
            content,
            base_filename
        )
        return path
    
    def _create_docx_document_sync(
        self,
        content: Dict[str, Any],
        base_filename: str
    ) -> Path:
        """Synchronous DOCX creation (runs in thread pool)"""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        doc = Document()
        
        # Set document properties
        metadata = content.get("document_metadata", {})
        doc_content = content.get("document_content", {})
        
        full_title = metadata.get("title", "Compliance Document")
        # Word document properties have a 255 character limit
        if len(full_title) > 255:
            truncated_title = full_title[:252] + "..."
        else:
            truncated_title = full_title
        
        doc.core_properties.title = truncated_title
        doc.core_properties.author = 'Compliance Agent'
        doc.core_properties.subject = metadata.get("document_type", "Compliance")
        doc.core_properties.created = datetime.utcnow()
        
        # Configure default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title page
        for _ in range(8):  # Add spacing
            doc.add_paragraph()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(metadata.get("title", "Compliance Document"))
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(24)
        title_run.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(f"Document Type: {metadata.get('document_type', 'N/A')}")
        meta_run.font.name = 'Calibri'
        meta_run.font.size = Pt(12)
        
        version_para = doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = version_para.add_run(f"Version: {metadata.get('version', '1.0')}")
        version_run.font.name = 'Calibri'
        version_run.font.size = Pt(12)
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.utcnow().strftime('%B %d, %Y'))
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(12)
        
        # Add page break
        doc.add_page_break()
        
        # Add Table of Contents
        doc.add_heading('Table of Contents', 1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        fldChar.set(qn('w:dirty'), 'true')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        run._element.append(instrText)
        run._element.append(fldChar2)
        doc.add_paragraph("")
        
        doc.add_page_break()
        
        # Add executive summary if present
        if "executive_summary" in doc_content:
            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(doc_content["executive_summary"])
            doc.add_paragraph()
        
        # Add main sections
        for section in doc_content.get("main_sections", []):
            # Add section heading
            heading = doc.add_heading(section.get("section_title", "Section"), 1)
            heading.style.font.name = 'Calibri'
            
            # Add section content
            section_text = section.get("section_content", "")
            if section_text:
                # Handle long text by splitting into paragraphs
                for para_text in section_text.split('\n\n'):
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.style.font.name = 'Calibri'
            
            # Add subsections
            for subsection in section.get("subsections", []):
                sub_heading = doc.add_heading(subsection.get("subsection_title", "Subsection"), 2)
                sub_heading.style.font.name = 'Calibri'
                
                subsection_text = subsection.get("subsection_content", "")
                if subsection_text:
                    para = doc.add_paragraph(subsection_text)
                    para.style.font.name = 'Calibri'
            
            doc.add_paragraph()  # Add spacing
        
        # Add appendices if present
        if "appendices" in doc_content and doc_content["appendices"]:
            doc.add_page_break()
            doc.add_heading("Appendices", 1)
            for appendix in doc_content["appendices"]:
                app_heading = doc.add_heading(appendix.get("appendix_title", "Appendix"), 2)
                app_heading.style.font.name = 'Calibri'
                app_para = doc.add_paragraph(appendix.get("appendix_content", ""))
                app_para.style.font.name = 'Calibri'
        
        # Add footer if present
        if "footer_content" in doc_content:
            doc.add_page_break()
            footer_para = doc.add_paragraph(doc_content["footer_content"])
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.name = 'Calibri'
        
        # Save document
        output_path = self.output_dir / f"{base_filename}.docx"
        doc.save(str(output_path))
        
        return output_path
    
    async def create_pdf_document(
        self,
        content: Dict[str, Any],
        base_filename: str
    ) -> Path:
        """Create a PDF document from content data asynchronously"""
        
        if not PDF_SUPPORT:
            logger.error("PDF generation requested but ReportLab is not installed")
            # Fall back to DOCX
            return await self.create_docx_document(content, base_filename)
        
        # Run in thread pool to avoid blocking
        path = await asyncio.to_thread(
            self._create_pdf_document_sync,
            content,
            base_filename
        )
        return path
    
    def _create_pdf_document_sync(
        self,
        content: Dict[str, Any],
        base_filename: str
    ) -> Path:
        """Synchronous PDF creation (runs in thread pool)"""
        
        try:
            output_path = self.output_dir / f"{base_filename}.pdf"
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=24,
                textColor=colors.HexColor('#2D3436'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading1_style = ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#2D3436'),
                spaceAfter=12,
                spaceBefore=12,
                leftIndent=0
            )
            
            heading2_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2D3436'),
                spaceAfter=10,
                spaceBefore=10,
                leftIndent=0
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceAfter=12,
                leading=14
            )
            
            # Get document properties
            metadata = content.get("document_metadata", {})
            doc_content = content.get("document_content", {})
            
            # Add title page
            elements.append(Spacer(1, 2*inch))
            elements.append(Paragraph(metadata.get("title", "Compliance Document"), title_style))
            elements.append(Spacer(1, 0.5*inch))
            
            # Add metadata
            elements.append(Paragraph(f"Document Type: {metadata.get('document_type', 'N/A')}", styles['Normal']))
            elements.append(Paragraph(f"Version: {metadata.get('version', '1.0')}", styles['Normal']))
            elements.append(Paragraph(datetime.utcnow().strftime('%B %d, %Y'), styles['Normal']))
            elements.append(PageBreak())
            
            # Add executive summary if present
            if "executive_summary" in doc_content:
                elements.append(Paragraph("Executive Summary", heading1_style))
                exec_text = self._clean_text_for_pdf(doc_content["executive_summary"])
                elements.append(Paragraph(exec_text, body_style))
                elements.append(Spacer(1, 12))
            
            # Add main sections
            for section in doc_content.get("main_sections", []):
                elements.append(Paragraph(section.get("section_title", "Section"), heading1_style))
                
                section_text = section.get("section_content", "")
                if section_text:
                    # Handle long text by splitting into paragraphs
                    for para_text in section_text.split('\n\n'):
                        if para_text.strip():
                            clean_text = self._clean_text_for_pdf(para_text.strip())
                            elements.append(Paragraph(clean_text, body_style))
                            elements.append(Spacer(1, 6))
                
                # Add subsections
                for subsection in section.get("subsections", []):
                    elements.append(Paragraph(subsection.get("subsection_title", "Subsection"), heading2_style))
                    subsection_text = subsection.get("subsection_content", "")
                    if subsection_text:
                        clean_text = self._clean_text_for_pdf(subsection_text)
                        elements.append(Paragraph(clean_text, body_style))
                        elements.append(Spacer(1, 6))
            
            # Add appendices if present
            if "appendices" in doc_content and doc_content["appendices"]:
                elements.append(PageBreak())
                elements.append(Paragraph("Appendices", heading1_style))
                for appendix in doc_content["appendices"]:
                    elements.append(Paragraph(appendix.get("appendix_title", "Appendix"), heading2_style))
                    app_text = self._clean_text_for_pdf(appendix.get("appendix_content", ""))
                    elements.append(Paragraph(app_text, body_style))
                    elements.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(elements)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating PDF document: {e}")
            # Fall back to DOCX
            return self._create_docx_document_sync(content, base_filename)
    
    def _clean_text_for_pdf(self, text: str) -> str:
        """Clean text for PDF rendering - escape special characters"""
        # Escape special ReportLab characters
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        return text
    
    async def create_excel_document(
        self,
        content: Dict[str, Any],
        document_type: str,
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Create Excel document based on document type"""
        
        if not EXCEL_AVAILABLE:
            logger.error("Excel generation requested but openpyxl is not installed")
            # Fall back to JSON
            return await self.create_json_fallback(content, base_filename)
        
        # Route to appropriate Excel generator based on document type
        doc_type_lower = document_type.lower()
        
        if "ropa" in doc_type_lower or "records_of_processing" in doc_type_lower:
            return await self.create_ropa_excel(content, framework, company_name, base_filename)
        elif "dpia" in doc_type_lower:
            return await self.create_dpia_excel(content, framework, company_name, base_filename)
        elif "checklist" in doc_type_lower:
            return await self.create_compliance_checklist_excel(content, framework, company_name, base_filename)
        elif "vendor" in doc_type_lower:
            return await self.create_vendor_assessment_excel(content, framework, company_name, base_filename)
        elif "training" in doc_type_lower:
            return await self.create_training_materials_excel(content, framework, company_name, base_filename)
        else:
            # Default to checklist format for unknown types
            return await self.create_compliance_checklist_excel(content, framework, company_name, base_filename)
    
    async def create_ropa_excel(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Create ROPA Excel document"""
        
        path = await asyncio.to_thread(
            self._create_ropa_excel_sync,
            content,
            framework,
            company_name,
            base_filename
        )
        return path
    
    def _create_ropa_excel_sync(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Synchronous ROPA Excel creation"""
        
        wb = Workbook()
        ws = wb.active
        ws.title = "ROPA Register"
        
        # Define styles
        header_font = Font(bold=True, color="FFFFFF", size=12)
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        # Add company header
        ws.merge_cells('A1:J1')
        ws['A1'] = f"{company_name} - Records of Processing Activities (ROPA)"
        ws['A1'].font = Font(bold=True, size=16)
        ws['A1'].alignment = Alignment(horizontal="center")
        
        ws.merge_cells('A2:J2')
        ws['A2'] = f"Framework: {framework.upper()} | Generated: {datetime.utcnow().strftime('%Y-%m-%d')}"
        ws['A2'].alignment = Alignment(horizontal="center")
        
        # ROPA Headers
        headers = [
            "Activity ID",
            "Processing Activity",
            "Purpose of Processing",
            "Legal Basis",
            "Categories of Data",
            "Data Subjects",
            "Recipients",
            "Retention Period",
            "Security Measures",
            "Third Country Transfers"
        ]
        
        # Add headers
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Extract ROPA data from content
        doc_content = content.get("document_content", {})
        main_sections = doc_content.get("main_sections", [])
        
        # Sample ROPA entries
        ropa_entries = [
            {
                "id": "PROC-001",
                "activity": "Customer Order Processing",
                "purpose": "Fulfillment of customer orders and contracts",
                "legal_basis": "Contract performance",
                "data_categories": "Name, Address, Contact details, Order history",
                "data_subjects": "Customers",
                "recipients": "Fulfillment partners, Payment processors",
                "retention": "7 years for tax purposes",
                "security": "Encryption at rest and in transit, Access controls",
                "transfers": "None"
            },
            {
                "id": "PROC-002",
                "activity": "Marketing Communications",
                "purpose": "Direct marketing and customer engagement",
                "legal_basis": "Consent",
                "data_categories": "Email, Name, Preferences",
                "data_subjects": "Customers, Prospects",
                "recipients": "Marketing platforms",
                "retention": "Until consent withdrawn",
                "security": "Secure APIs, Encryption",
                "transfers": "EU adequacy decision areas only"
            }
        ]
        
        # Add data rows
        for row_idx, entry in enumerate(ropa_entries, start=5):
            ws.cell(row=row_idx, column=1, value=entry.get("id"))
            ws.cell(row=row_idx, column=2, value=entry.get("activity"))
            ws.cell(row=row_idx, column=3, value=entry.get("purpose"))
            ws.cell(row=row_idx, column=4, value=entry.get("legal_basis"))
            ws.cell(row=row_idx, column=5, value=entry.get("data_categories"))
            ws.cell(row=row_idx, column=6, value=entry.get("data_subjects"))
            ws.cell(row=row_idx, column=7, value=entry.get("recipients"))
            ws.cell(row=row_idx, column=8, value=entry.get("retention"))
            ws.cell(row=row_idx, column=9, value=entry.get("security"))
            ws.cell(row=row_idx, column=10, value=entry.get("transfers"))
        
        # Adjust column widths
        column_widths = [12, 25, 25, 20, 25, 20, 25, 20, 25, 20]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width
        
        # Add borders
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in ws.iter_rows(min_row=4, max_row=len(ropa_entries)+4, min_col=1, max_col=10):
            for cell in row:
                cell.border = thin_border
                if cell.row > 4:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
        
        # Save workbook
        output_path = self.output_dir / f"{base_filename}.xlsx"
        wb.save(str(output_path))
        
        return output_path
    
    async def create_dpia_excel(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Create DPIA Excel document"""
        
        path = await asyncio.to_thread(
            self._create_dpia_excel_sync,
            content,
            framework,
            company_name,
            base_filename
        )
        return path
    
    def _create_dpia_excel_sync(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Synchronous DPIA Excel creation"""
        
        wb = Workbook()
        
        # Create multiple sheets for DPIA sections
        sheets = {
            "Overview": wb.active,
            "Risk Assessment": wb.create_sheet("Risk Assessment"),
            "Mitigation Measures": wb.create_sheet("Mitigation Measures"),
            "Compliance Checklist": wb.create_sheet("Compliance Checklist")
        }
        
        # Overview Sheet
        ws = sheets["Overview"]
        ws.title = "Overview"
        
        # Add headers and content
        overview_data = [
            ["DPIA - Data Protection Impact Assessment"],
            [f"Organization: {company_name}"],
            [f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}"],
            [f"Framework: {framework.upper()}"],
            [""],
            ["Section", "Status", "Comments"],
            ["1. Processing Description", "Complete", "Detailed in main document"],
            ["2. Necessity Assessment", "Complete", "Justified per legal basis"],
            ["3. Risk Identification", "Complete", "See Risk Assessment tab"],
            ["4. Mitigation Measures", "Complete", "See Mitigation tab"],
            ["5. Consultation", "Pending", "If high risk remains"],
            ["6. Review Schedule", "Annual", "Next review in 12 months"]
        ]
        
        for row_idx, row_data in enumerate(overview_data, 1):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Risk Assessment Sheet
        ws = sheets["Risk Assessment"]
        risk_headers = ["Risk ID", "Risk Description", "Likelihood", "Impact", "Risk Level", "Mitigation Required"]
        
        for col, header in enumerate(risk_headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
        
        # Sample risk data
        risks = [
            ["RISK-001", "Unauthorized access to personal data", "Medium", "High", "High", "Yes"],
            ["RISK-002", "Data breach during transfer", "Low", "High", "Medium", "Yes"],
            ["RISK-003", "Excessive data retention", "Low", "Medium", "Low", "Yes"],
            ["RISK-004", "Third-party processor breach", "Medium", "High", "High", "Yes"]
        ]
        
        for row_idx, risk_data in enumerate(risks, 2):
            for col_idx, value in enumerate(risk_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Adjust column widths for all sheets
        for sheet_name, sheet in sheets.items():
            for column in sheet.columns:
                max_length = 0
                column = [cell for cell in column]
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                if adjusted_width > 0:
                    sheet.column_dimensions[column[0].column_letter].width = adjusted_width
        
        # Save workbook
        output_path = self.output_dir / f"{base_filename}.xlsx"
        wb.save(str(output_path))
        
        return output_path
    
    async def create_compliance_checklist_excel(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Create Compliance Checklist Excel document"""
        
        path = await asyncio.to_thread(
            self._create_compliance_checklist_excel_sync,
            content,
            framework,
            company_name,
            base_filename
        )
        return path
    
    def _create_compliance_checklist_excel_sync(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Synchronous Compliance Checklist Excel creation"""
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Compliance Checklist"
        
        # Add header
        ws['A1'] = f"{company_name} - {framework.upper()} Compliance Checklist"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:F1')
        
        # Checklist headers
        headers = ["Category", "Requirement", "Status", "Evidence", "Owner", "Due Date"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
        
        # Sample checklist items (could extract from content)
        checklist_items = [
            ["Data Protection", "Privacy Policy Published", "Complete", "Website URL", "Legal Team", "Completed"],
            ["Data Protection", "DPIA Conducted", "In Progress", "DPIA Document", "DPO", "2024-Q1"],
            ["Security", "Access Controls Implemented", "Complete", "IAM Policy", "IT Security", "Completed"],
            ["Security", "Encryption at Rest", "Complete", "Technical Spec", "IT Security", "Completed"],
            ["Governance", "DPO Appointed", "Complete", "Appointment Letter", "Board", "Completed"],
            ["Training", "Staff Training Completed", "In Progress", "Training Records", "HR", "2024-Q1"]
        ]
        
        for row_idx, item in enumerate(checklist_items, 4):
            for col_idx, value in enumerate(item, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Auto-adjust columns
        for column in ws.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 40)
            if adjusted_width > 0:
                ws.column_dimensions[column[0].column_letter].width = adjusted_width
        
        # Save workbook
        output_path = self.output_dir / f"{base_filename}.xlsx"
        wb.save(str(output_path))
        
        return output_path
    
    async def create_vendor_assessment_excel(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Create Vendor Assessment Excel document"""
        
        path = await asyncio.to_thread(
            self._create_vendor_assessment_excel_sync,
            content,
            framework,
            company_name,
            base_filename
        )
        return path
    
    def _create_vendor_assessment_excel_sync(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Synchronous Vendor Assessment Excel creation"""
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Vendor Assessment"
        
        # Add assessment criteria
        ws['A1'] = f"{company_name} - Vendor/Processor Assessment"
        ws['A1'].font = Font(bold=True, size=14)
        
        headers = ["Vendor Name", "Service Type", "Data Processed", "Risk Level", "Contract Status", "Assessment Date", "Next Review"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=3, column=col, value=header).font = Font(bold=True)
        
        # Sample vendor data
        vendors = [
            ["Cloud Provider A", "Infrastructure", "All customer data", "High", "Signed", "2024-01-01", "2024-07-01"],
            ["Payment Processor B", "Payments", "Payment data", "High", "Signed", "2024-01-01", "2024-07-01"],
            ["Email Service C", "Communications", "Email addresses", "Medium", "Under Review", "2024-02-01", "2024-08-01"]
        ]
        
        for row_idx, vendor in enumerate(vendors, 4):
            for col_idx, value in enumerate(vendor, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Auto-adjust columns
        for column in ws.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 40)
            if adjusted_width > 0:
                ws.column_dimensions[column[0].column_letter].width = adjusted_width
        
        # Save workbook
        output_path = self.output_dir / f"{base_filename}.xlsx"
        wb.save(str(output_path))
        
        return output_path
    
    async def create_training_materials_excel(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Create Training Materials Excel document"""
        
        path = await asyncio.to_thread(
            self._create_training_materials_excel_sync,
            content,
            framework,
            company_name,
            base_filename
        )
        return path
    
    def _create_training_materials_excel_sync(
        self,
        content: Dict[str, Any],
        framework: str,
        company_name: str,
        base_filename: str
    ) -> Path:
        """Synchronous Training Materials Excel creation"""
        
        wb = Workbook()
        
        # Create sheets for different training modules
        sheets = {
            "Overview": wb.active,
            "Module Plan": wb.create_sheet("Module Plan"),
            "Quiz Questions": wb.create_sheet("Quiz Questions"),
            "Resources": wb.create_sheet("Resources")
        }
        
        # Overview Sheet
        ws = sheets["Overview"]
        ws.title = "Overview"
        
        overview_data = [
            [f"{company_name} - {framework.upper()} Training Program"],
            [f"Generated: {datetime.utcnow().strftime('%Y-%m-%d')}"],
            [""],
            ["Module", "Duration", "Target Audience", "Status"],
            ["Introduction to Compliance", "1 hour", "All Staff", "Available"],
            ["Data Protection Basics", "2 hours", "All Staff", "Available"],
            ["Security Best Practices", "1.5 hours", "Technical Staff", "In Development"],
            ["Incident Response", "1 hour", "Management", "In Development"]
        ]
        
        for row_idx, row_data in enumerate(overview_data, 1):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Auto-adjust columns for all sheets
        for sheet_name, sheet in sheets.items():
            for column in sheet.columns:
                max_length = 0
                column = [cell for cell in column]
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 40)
                if adjusted_width > 0:
                    sheet.column_dimensions[column[0].column_letter].width = adjusted_width
        
        # Save workbook
        output_path = self.output_dir / f"{base_filename}.xlsx"
        wb.save(str(output_path))
        
        return output_path
    
    async def create_json_fallback(
        self,
        content: Dict[str, Any],
        base_filename: str
    ) -> Path:
        """Create JSON document as fallback"""
        
        import json
        
        output_path = self.output_dir / f"{base_filename}.json"
        
        # Write formatted JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, indent=2, ensure_ascii=False, default=str)
        
        return output_path

# Create singleton instance
compliance_document_service = ComplianceDocumentService()
