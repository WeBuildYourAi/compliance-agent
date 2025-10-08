"""
Enhanced Marketing Document Generation Workflow
Handles multi-document generation with quality validation and file outputs
"""
import sys
import os
from pathlib import Path

# Add the src directory to Python path for imports
current_dir = Path(__file__).parent
if str(current_dir) not in sys.path:
    sys.path.insert(0, str(current_dir))

from typing import Dict, Any, List, Optional, TypedDict, Literal, Annotated
from enum import Enum
from langgraph.graph import StateGraph, END
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from datetime import datetime
import logging
import uuid
import json
import asyncio
from concurrent.futures import ThreadPoolExecutor
import base64
from io import BytesIO

# Configure logging FIRST before using logger
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Document generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX generation will be limited.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("reportlab not installed. PDF generation will be limited.")

# Import utilities
from llm_utils import llm_manager
from config import config
from storage_utils import storage_manager, cosmos_manager

class DocumentStatus(str, Enum):
    PENDING = "pending"
    GENERATING = "generating"
    REVIEWING = "reviewing"
    VALIDATING = "validating"
    COMPLETED = "completed"
    FAILED = "failed"
    REQUIRES_REVISION = "requires_revision"

class ValidationStatus(str, Enum):
    NOT_VALIDATED = "not_validated"
    PASSED = "passed"
    FAILED = "failed"
    PARTIAL = "partial"

# Enhanced State Definition for Marketing Document Generation
class MarketingDocumentState(TypedDict):
    # Core Input from orchestrator
    user_prompt: str
    project_brief: Dict[str, Any]  # Research synthesis from previous agent
    deliverable_blueprint: List[str]  # List of documents to create
    user_responses: Optional[Dict[str, Any]]  # User's answers to questions
    interaction_questions: Optional[List[Dict[str, Any]]]  # Questions that were asked
    request_id: str
    correlation_id: str
    
    # Session Management
    user_id: Optional[str]
    session_id: Optional[str]
    
    # Document Planning
    document_plans: List[Dict[str, Any]]  # Detailed plan for each document
    document_dependencies: Dict[str, List[str]]  # Dependencies between documents
    
    # Document Generation & Storage
    document_contents: Dict[str, Dict[str, Any]]  # Generated content for each document
    document_files: Dict[str, str]  # File paths for generated documents
    document_urls: List[str]  # Public URLs for accessing documents
    document_status: Dict[str, DocumentStatus]
    
    # Validation & Quality
    validation_results: Dict[str, Dict[str, Any]]  # Validation results per document
    cross_document_validation: Dict[str, Any]  # Consistency checks across documents
    quality_scores: Dict[str, float]  # Quality score for each document
    overall_validation_status: ValidationStatus
    
    # Success Criteria Tracking
    success_criteria: List[str]  # From project brief
    criteria_validation: Dict[str, bool]  # Whether each criterion is met
    
    # Workflow State
    status: str
    current_stage: str
    error_message: Optional[str]
    retry_count: int
    
    # Output
    final_deliverables: Dict[str, Any]  # Final package with all documents
    executive_summary: Dict[str, Any]
    document_manifest: List[Dict[str, Any]]  # List of all documents with metadata
    
    # Message History
    messages: Annotated[List[BaseMessage], "Message history for conversation tracking"]
    
    # Metadata
    created_at: datetime
    updated_at: datetime
    processing_duration: Optional[float]

def ensure_state_initialization(state: MarketingDocumentState) -> MarketingDocumentState:
    """Ensure all required state fields are initialized"""
    # Core structures
    if "messages" not in state or state["messages"] is None:
        state["messages"] = []
    if "document_plans" not in state or state["document_plans"] is None:
        state["document_plans"] = []
    if "document_contents" not in state or state["document_contents"] is None:
        state["document_contents"] = {}
    if "document_files" not in state or state["document_files"] is None:
        state["document_files"] = {}
    if "document_urls" not in state or state["document_urls"] is None:
        state["document_urls"] = []
    if "document_status" not in state or state["document_status"] is None:
        state["document_status"] = {}
    if "validation_results" not in state or state["validation_results"] is None:
        state["validation_results"] = {}
    if "quality_scores" not in state or state["quality_scores"] is None:
        state["quality_scores"] = {}
    if "criteria_validation" not in state or state["criteria_validation"] is None:
        state["criteria_validation"] = {}
    if "document_dependencies" not in state or state["document_dependencies"] is None:
        state["document_dependencies"] = {}
    if "cross_document_validation" not in state or state["cross_document_validation"] is None:
        state["cross_document_validation"] = {}
    if "final_deliverables" not in state or state["final_deliverables"] is None:
        state["final_deliverables"] = {}
    if "executive_summary" not in state or state["executive_summary"] is None:
        state["executive_summary"] = {}
    if "document_manifest" not in state or state["document_manifest"] is None:
        state["document_manifest"] = []
    
    # Extract success criteria from project brief if not set
    if "success_criteria" not in state or state["success_criteria"] is None:
        project_brief = state.get("project_brief", {})
        state["success_criteria"] = project_brief.get("success_criteria", [])
    
    return state

def handle_node_error(state: MarketingDocumentState, error: Exception, node_name: str) -> MarketingDocumentState:
    """Centralized error handling for all nodes"""
    logger.error(f"Error in {node_name}: {error}", exc_info=True)
    
    state = ensure_state_initialization(state)
    
    state["error_message"] = f"{node_name} failed: {str(error)}"
    state["retry_count"] = state.get("retry_count", 0) + 1
    state["updated_at"] = datetime.utcnow()
    
    state["messages"].append(AIMessage(
        content=f"Processing error in {node_name}. Retry attempt {state['retry_count']}/3."
    ))
    
    return state

async def analyze_document_requirements(state: MarketingDocumentState) -> MarketingDocumentState:
    """Analyze deliverable blueprint and create detailed document plans"""
    
    try:
        logger.info(f"Analyzing document requirements for request: {state.get('request_id')}")
        
        state = ensure_state_initialization(state)
        
        state["status"] = "analyzing_requirements"
        state["current_stage"] = "requirements_analysis"
        state["updated_at"] = datetime.utcnow()
        
        deliverable_blueprint = state.get("deliverable_blueprint", [])
        project_brief = state.get("project_brief", {})
        user_responses = state.get("user_responses", {})
        
        if not deliverable_blueprint:
            raise ValueError("No deliverable blueprint provided")
        
        llm = llm_manager.get_standard_llm()
        
        # Create detailed plan for each document
        document_plans = []
        
        for i, deliverable in enumerate(deliverable_blueprint):
            planning_prompt = f"""
            You are a marketing document specialist. Create a detailed execution plan for this deliverable.
            
            DELIVERABLE: {deliverable}
            
            PROJECT CONTEXT:
            Research Summary: {project_brief.get('research_summary', 'Not provided')}
            Success Criteria: {json.dumps(state.get('success_criteria', []), indent=2)}
            User Context: {json.dumps(user_responses, indent=2)}
            
            Create a detailed document plan in JSON format:
            {{
                "document_id": "doc_{i+1:03d}",
                "title": "Document Title",
                "type": "brief|schema|checklist|guide|calendar|matrix|specification|dashboard|playbook",
                "description": "What this document accomplishes",
                "target_audience": "Who will use this document",
                "format": "docx|pdf|xlsx|json|yaml",
                "estimated_length": "pages or sections",
                "key_sections": [
                    {{
                        "section_name": "Section Name",
                        "section_description": "What this section covers",
                        "content_type": "narrative|technical|visual|data",
                        "estimated_length": "paragraphs/pages"
                    }}
                ],
                "dependencies": ["List of other document IDs this depends on"],
                "validation_criteria": [
                    "Specific criteria to validate this document"
                ],
                "quality_requirements": [
                    "Quality standards this document must meet"
                ]
            }}
            """
            
            plan_result = await llm_manager.safe_llm_query(llm, planning_prompt, parse_json=True)
            
            if "error" not in plan_result:
                plan_result["deliverable_description"] = deliverable
                document_plans.append(plan_result)
            else:
                # Fallback plan
                document_plans.append({
                    "document_id": f"doc_{i+1:03d}",
                    "title": deliverable[:50],
                    "type": "document",
                    "deliverable_description": deliverable,
                    "format": "docx",
                    "key_sections": [{"section_name": "Main Content"}]
                })
        
        # Identify dependencies between documents
        document_dependencies = {}
        for plan in document_plans:
            doc_id = plan["document_id"]
            deps = plan.get("dependencies", [])
            document_dependencies[doc_id] = deps
        
        # Update state
        state["document_plans"] = document_plans
        state["document_dependencies"] = document_dependencies
        
        # Initialize document status
        for plan in document_plans:
            state["document_status"][plan["document_id"]] = DocumentStatus.PENDING
        
        state["messages"].append(AIMessage(
            content=f"Document planning complete: {len(document_plans)} documents planned with dependencies mapped."
        ))
        
        return state
        
    except Exception as e:
        return handle_node_error(state, e, "analyze_document_requirements")

async def generate_document_content(doc_plan: Dict[str, Any], context: Dict[str, Any]) -> Dict[str, Any]:
    """Generate content for a single document"""
    
    doc_id = doc_plan.get("document_id")
    
    try:
        logger.info(f"Generating content for: {doc_id} - {doc_plan.get('title')}")
        
        llm = llm_manager.get_standard_llm()
        
        generation_prompt = f"""
        You are a marketing document specialist. Generate comprehensive content for this document.
        
        DOCUMENT PLAN:
        {json.dumps(doc_plan, indent=2)}
        
        PROJECT CONTEXT:
        Research Summary: {context.get('research_summary', 'Not provided')}
        Success Criteria: {json.dumps(context.get('success_criteria', []), indent=2)}
        User Context: {json.dumps(context.get('user_responses', {}), indent=2)}
        
        Generate complete document content in JSON format:
        {{
            "document_metadata": {{
                "title": "{doc_plan.get('title')}",
                "type": "{doc_plan.get('type')}",
                "version": "1.0",
                "created_date": "{datetime.utcnow().isoformat()}",
                "target_audience": "{doc_plan.get('target_audience', 'General')}"
            }},
            "executive_summary": "Brief overview of this document",
            "sections": [
                {{
                    "section_title": "Section Title",
                    "content": "Detailed content for this section",
                    "subsections": [
                        {{
                            "title": "Subsection Title",
                            "content": "Subsection content"
                        }}
                    ]
                }}
            ],
            "key_takeaways": [
                "Key point 1",
                "Key point 2"
            ],
            "next_steps": [
                "Action item 1",
                "Action item 2"
            ],
            "appendices": [
                {{
                    "title": "Appendix Title",
                    "content": "Supporting information"
                }}
            ]
        }}
        
        REQUIREMENTS:
        1. Create professional, actionable content
        2. Ensure alignment with the deliverable description
        3. Include specific examples and implementation details
        4. Make content immediately usable by the target audience
        5. Follow marketing best practices
        """
        
        content_result = await llm_manager.safe_llm_query(llm, generation_prompt, parse_json=True)
        
        if "error" not in content_result:
            return {
                "document_id": doc_id,
                "status": "generated",
                "content": content_result,
                "success": True
            }
        else:
            return {
                "document_id": doc_id,
                "status": "failed",
                "error": content_result.get("error"),
                "success": False
            }
    
    except Exception as e:
        logger.error(f"Content generation failed for {doc_id}: {e}")
        return {
            "document_id": doc_id,
            "status": "failed",
            "error": str(e),
            "success": False
        }

async def generate_all_documents(state: MarketingDocumentState) -> MarketingDocumentState:
    """Generate content for all planned documents"""
    
    try:
        logger.info("Starting document generation phase")
        
        state = ensure_state_initialization(state)
        
        state["status"] = "generating_documents"
        state["current_stage"] = "document_generation"
        state["updated_at"] = datetime.utcnow()
        
        document_plans = state.get("document_plans", [])
        project_brief = state.get("project_brief", {})
        
        # Prepare generation context
        generation_context = {
            "research_summary": project_brief.get("research_summary", ""),
            "success_criteria": state.get("success_criteria", []),
            "user_responses": state.get("user_responses", {}),
            "implementation_approach": project_brief.get("implementation_approach", "")
        }
        
        # Check for dependencies and generate in order
        generated_docs = {}
        pending_docs = document_plans.copy()
        max_iterations = len(document_plans) * 2  # Prevent infinite loops
        iteration = 0
        
        while pending_docs and iteration < max_iterations:
            iteration += 1
            docs_to_generate = []
            
            for doc_plan in pending_docs:
                doc_id = doc_plan["document_id"]
                deps = state["document_dependencies"].get(doc_id, [])
                
                # Check if all dependencies are satisfied
                if all(dep in generated_docs for dep in deps):
                    docs_to_generate.append(doc_plan)
            
            if not docs_to_generate:
                # No documents can be generated (circular dependency or error)
                logger.warning("Cannot resolve document dependencies")
                break
            
            # Generate documents in parallel batch
            tasks = []
            for doc_plan in docs_to_generate:
                # Add previously generated docs to context for reference
                context = generation_context.copy()
                context["previous_documents"] = {
                    dep_id: generated_docs[dep_id].get("content", {}).get("executive_summary", "")
                    for dep_id in state["document_dependencies"].get(doc_plan["document_id"], [])
                    if dep_id in generated_docs
                }
                
                tasks.append(generate_document_content(doc_plan, context))
            
            # Execute batch
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process results
            for doc_plan, result in zip(docs_to_generate, results):
                doc_id = doc_plan["document_id"]
                
                if isinstance(result, Exception):
                    state["document_status"][doc_id] = DocumentStatus.FAILED
                    state["document_contents"][doc_id] = {"error": str(result)}
                else:
                    if result.get("success"):
                        state["document_status"][doc_id] = DocumentStatus.REVIEWING
                        state["document_contents"][doc_id] = result.get("content", {})
                        generated_docs[doc_id] = result
                    else:
                        state["document_status"][doc_id] = DocumentStatus.FAILED
                        state["document_contents"][doc_id] = {"error": result.get("error")}
                
                # Remove from pending
                pending_docs = [d for d in pending_docs if d["document_id"] != doc_id]
        
        # Count results
        successful = len([s for s in state["document_status"].values() if s == DocumentStatus.REVIEWING])
        total = len(document_plans)
        
        state["messages"].append(AIMessage(
            content=f"Document generation complete: {successful}/{total} documents generated successfully."
        ))
        
        return state
        
    except Exception as e:
        return handle_node_error(state, e, "generate_all_documents")

async def validate_individual_documents(state: MarketingDocumentState) -> MarketingDocumentState:
    """Validate each document against requirements and quality standards"""
    
    try:
        logger.info("Starting individual document validation")
        
        state = ensure_state_initialization(state)
        
        state["status"] = "validating_documents"
        state["current_stage"] = "individual_validation"
        state["updated_at"] = datetime.utcnow()
        
        llm = llm_manager.get_mini_llm()  # Use mini model for validation
        
        for doc_id, content in state["document_contents"].items():
            if state["document_status"][doc_id] != DocumentStatus.REVIEWING:
                continue
            
            # Find the original plan for this document
            doc_plan = next((p for p in state["document_plans"] if p["document_id"] == doc_id), None)
            if not doc_plan:
                continue
            
            validation_prompt = f"""
            Validate this document against its requirements and quality standards.
            
            DOCUMENT PLAN:
            {json.dumps(doc_plan, indent=2)}
            
            GENERATED CONTENT:
            {json.dumps(content, indent=2)[:3000]}...  # Truncate for context window
            
            SUCCESS CRITERIA TO CHECK:
            {json.dumps(state.get('success_criteria', []), indent=2)}
            
            Provide validation results in JSON format:
            {{
                "overall_pass": true/false,
                "quality_score": 0.0-1.0,
                "criteria_met": [
                    {{"criterion": "criterion text", "met": true/false, "notes": "explanation"}}
                ],
                "issues_found": [
                    {{"severity": "critical|major|minor", "issue": "description", "suggestion": "how to fix"}}
                ],
                "strengths": ["list of strong points"],
                "recommendations": ["list of improvement suggestions"]
            }}
            """
            
            validation_result = await llm_manager.safe_llm_query(llm, validation_prompt, parse_json=True)
            
            if "error" not in validation_result:
                state["validation_results"][doc_id] = validation_result
                state["quality_scores"][doc_id] = validation_result.get("quality_score", 0.5)
                
                # Update status based on validation
                if validation_result.get("overall_pass", False):
                    state["document_status"][doc_id] = DocumentStatus.VALIDATED
                else:
                    critical_issues = [i for i in validation_result.get("issues_found", []) 
                                     if i.get("severity") == "critical"]
                    if critical_issues:
                        state["document_status"][doc_id] = DocumentStatus.REQUIRES_REVISION
                    else:
                        state["document_status"][doc_id] = DocumentStatus.VALIDATED
            else:
                state["validation_results"][doc_id] = {"error": "Validation failed"}
                state["quality_scores"][doc_id] = 0.5
        
        state["messages"].append(AIMessage(
            content=f"Individual validation complete. Average quality score: {sum(state['quality_scores'].values())/len(state['quality_scores']):.2f}"
        ))
        
        return state
        
    except Exception as e:
        return handle_node_error(state, e, "validate_individual_documents")

async def validate_cross_document_consistency(state: MarketingDocumentState) -> MarketingDocumentState:
    """Validate consistency and alignment across all documents"""
    
    try:
        logger.info("Starting cross-document validation")
        
        state = ensure_state_initialization(state)
        
        state["status"] = "validating_consistency"
        state["current_stage"] = "cross_validation"
        state["updated_at"] = datetime.utcnow()
        
        llm = llm_manager.get_standard_llm()
        
        # Prepare summary of all documents for consistency check
        doc_summaries = []
        for doc_id, content in state["document_contents"].items():
            if isinstance(content, dict) and "document_metadata" in content:
                doc_summaries.append({
                    "document_id": doc_id,
                    "title": content.get("document_metadata", {}).get("title", "Unknown"),
                    "executive_summary": content.get("executive_summary", ""),
                    "key_takeaways": content.get("key_takeaways", [])
                })
        
        consistency_prompt = f"""
        Validate consistency and alignment across all generated documents.
        
        DOCUMENT SUMMARIES:
        {json.dumps(doc_summaries, indent=2)}
        
        PROJECT SUCCESS CRITERIA:
        {json.dumps(state.get('success_criteria', []), indent=2)}
        
        CHECK FOR:
        1. Terminology consistency across documents
        2. Timeline and date alignment
        3. Metric and KPI consistency
        4. Brand voice and tone alignment
        5. Technical specification consistency
        6. Process and workflow alignment
        7. Coverage of all success criteria across documents
        
        Provide validation results in JSON format:
        {{
            "overall_consistency": true/false,
            "consistency_score": 0.0-1.0,
            "inconsistencies_found": [
                {{
                    "type": "terminology|timeline|metric|voice|technical|process",
                    "severity": "critical|major|minor",
                    "documents_affected": ["doc_id1", "doc_id2"],
                    "description": "What is inconsistent",
                    "resolution": "How to fix"
                }}
            ],
            "success_criteria_coverage": {{
                "all_criteria_covered": true/false,
                "coverage_by_criterion": [
                    {{"criterion": "text", "covered": true/false, "documents": ["doc_ids"]}}
                ]
            }},
            "alignment_strengths": ["What is well-aligned"],
            "recommendations": ["Overall improvement suggestions"]
        }}
        """
        
        consistency_result = await llm_manager.safe_llm_query(llm, consistency_prompt, parse_json=True)
        
        if "error" not in consistency_result:
            state["cross_document_validation"] = consistency_result
            
            # Update criteria validation
            criteria_coverage = consistency_result.get("success_criteria_coverage", {}).get("coverage_by_criterion", [])
            for item in criteria_coverage:
                criterion = item.get("criterion", "")
                state["criteria_validation"][criterion] = item.get("covered", False)
            
            # Determine overall validation status
            if consistency_result.get("overall_consistency", False) and consistency_result.get("consistency_score", 0) > 0.7:
                state["overall_validation_status"] = ValidationStatus.PASSED
            else:
                critical_issues = [i for i in consistency_result.get("inconsistencies_found", [])
                                 if i.get("severity") == "critical"]
                if critical_issues:
                    state["overall_validation_status"] = ValidationStatus.FAILED
                else:
                    state["overall_validation_status"] = ValidationStatus.PARTIAL
        else:
            state["cross_document_validation"] = {"error": "Cross-validation failed"}
            state["overall_validation_status"] = ValidationStatus.NOT_VALIDATED
        
        state["messages"].append(AIMessage(
            content=f"Cross-document validation complete. Consistency score: {consistency_result.get('consistency_score', 0):.2f}"
        ))
        
        return state
        
    except Exception as e:
        return handle_node_error(state, e, "validate_cross_document_consistency")

def create_docx_file(content: Dict[str, Any], title: str) -> bytes:
    """Create a DOCX file from document content"""
    
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX generation")
    
    doc = Document()
    
    # Add title
    doc.add_heading(title, 0)
    
    # Add metadata
    metadata = content.get("document_metadata", {})
    if metadata:
        doc.add_paragraph(f"Version: {metadata.get('version', '1.0')}")
        doc.add_paragraph(f"Created: {metadata.get('created_date', datetime.utcnow().strftime('%Y-%m-%d'))}")
        doc.add_paragraph(f"Target Audience: {metadata.get('target_audience', 'General')}")
        doc.add_paragraph()
    
    # Add executive summary
    if "executive_summary" in content:
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(content["executive_summary"])
        doc.add_paragraph()
    
    # Add main sections
    sections = content.get("sections", [])
    for section in sections:
        doc.add_heading(section.get("section_title", "Section"), 1)
        doc.add_paragraph(section.get("content", ""))
        
        # Add subsections
        for subsection in section.get("subsections", []):
            doc.add_heading(subsection.get("title", "Subsection"), 2)
            doc.add_paragraph(subsection.get("content", ""))
        
        doc.add_paragraph()
    
    # Add key takeaways
    if "key_takeaways" in content:
        doc.add_heading("Key Takeaways", 1)
        for takeaway in content["key_takeaways"]:
            doc.add_paragraph(f"• {takeaway}", style='List Bullet')
        doc.add_paragraph()
    
    # Add next steps
    if "next_steps" in content:
        doc.add_heading("Next Steps", 1)
        for i, step in enumerate(content["next_steps"], 1):
            doc.add_paragraph(f"{i}. {step}", style='List Number')
        doc.add_paragraph()
    
    # Add appendices
    for appendix in content.get("appendices", []):
        doc.add_page_break()
        doc.add_heading(appendix.get("title", "Appendix"), 1)
        doc.add_paragraph(appendix.get("content", ""))
    
    # Save to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

def create_pdf_file(content: Dict[str, Any], title: str) -> bytes:
    """Create a PDF file from document content"""
    
    if not PDF_AVAILABLE:
        raise ImportError("reportlab is required for PDF generation")
    
    pdf_bytes = BytesIO()
    doc = SimpleDocTemplate(pdf_bytes, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    # Metadata
    metadata = content.get("document_metadata", {})
    if metadata:
        for key, value in metadata.items():
            story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Executive Summary
    if "executive_summary" in content:
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        story.append(Paragraph(content["executive_summary"], styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Sections
    for section in content.get("sections", []):
        story.append(Paragraph(section.get("section_title", "Section"), styles['Heading2']))
        story.append(Paragraph(section.get("content", ""), styles['Normal']))
        story.append(Spacer(1, 6))
        
        for subsection in section.get("subsections", []):
            story.append(Paragraph(subsection.get("title", "Subsection"), styles['Heading3']))
            story.append(Paragraph(subsection.get("content", ""), styles['Normal']))
            story.append(Spacer(1, 6))
        
        story.append(Spacer(1, 12))
    
    # Key Takeaways
    if "key_takeaways" in content:
        story.append(Paragraph("Key Takeaways", styles['Heading2']))
        for takeaway in content["key_takeaways"]:
            story.append(Paragraph(f"• {takeaway}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Next Steps
    if "next_steps" in content:
        story.append(Paragraph("Next Steps", styles['Heading2']))
        for i, step in enumerate(content["next_steps"], 1):
            story.append(Paragraph(f"{i}. {step}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Appendices
    for appendix in content.get("appendices", []):
        story.append(PageBreak())
        story.append(Paragraph(appendix.get("title", "Appendix"), styles['Heading2']))
        story.append(Paragraph(appendix.get("content", ""), styles['Normal']))
    
    # Build PDF
    doc.build(story)
    pdf_bytes.seek(0)
    
    return pdf_bytes.getvalue()

async def create_document_files(state: MarketingDocumentState) -> MarketingDocumentState:
    """Create actual document files (DOCX, PDF, etc.) and save to storage"""
    
    try:
        logger.info("Creating document files")
        
        state = ensure_state_initialization(state)
        
        state["status"] = "creating_files"
        state["current_stage"] = "file_generation"
        state["updated_at"] = datetime.utcnow()
        
        output_dir = "/mnt/user-data/outputs"
        os.makedirs(output_dir, exist_ok=True)
        
        document_urls = []
        document_manifest = []
        
        for doc_plan in state["document_plans"]:
            doc_id = doc_plan["document_id"]
            content = state["document_contents"].get(doc_id, {})
            
            if not content or "error" in content:
                continue
            
            title = doc_plan.get("title", f"Document {doc_id}")
            format_type = doc_plan.get("format", "docx").lower()
            
            # Generate filename
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')[:50]
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            
            try:
                if format_type == "docx":
                    # Create DOCX file
                    filename = f"{safe_title}_{timestamp}.docx"
                    filepath = os.path.join(output_dir, filename)
                    
                    docx_bytes = create_docx_file(content, title)
                    with open(filepath, 'wb') as f:
                        f.write(docx_bytes)
                    
                    url = f"computer:///mnt/user-data/outputs/{filename}"
                    
                elif format_type == "pdf":
                    # Create PDF file
                    filename = f"{safe_title}_{timestamp}.pdf"
                    filepath = os.path.join(output_dir, filename)
                    
                    pdf_bytes = create_pdf_file(content, title)
                    with open(filepath, 'wb') as f:
                        f.write(pdf_bytes)
                    
                    url = f"computer:///mnt/user-data/outputs/{filename}"
                    
                else:
                    # Default to JSON for other formats
                    filename = f"{safe_title}_{timestamp}.json"
                    filepath = os.path.join(output_dir, filename)
                    
                    with open(filepath, 'w', encoding='utf-8') as f:
                        json.dump(content, f, indent=2, ensure_ascii=False)
                    
                    url = f"computer:///mnt/user-data/outputs/{filename}"
                
                # Store file info
                state["document_files"][doc_id] = filepath
                document_urls.append(url)
                
                # Add to manifest
                document_manifest.append({
                    "document_id": doc_id,
                    "title": title,
                    "filename": filename,
                    "format": format_type,
                    "url": url,
                    "quality_score": state["quality_scores"].get(doc_id, 0),
                    "validation_status": state["document_status"].get(doc_id, DocumentStatus.COMPLETED).value
                })
                
                # Update status
                state["document_status"][doc_id] = DocumentStatus.COMPLETED
                
                logger.info(f"Created file: {filename}")
                
            except Exception as e:
                logger.error(f"Failed to create file for {doc_id}: {e}")
                state["document_status"][doc_id] = DocumentStatus.FAILED
        
        # Update state with URLs and manifest
        state["document_urls"] = document_urls
        state["document_manifest"] = document_manifest
        
        state["messages"].append(AIMessage(
            content=f"File creation complete: {len(document_urls)} documents saved."
        ))
        
        return state
        
    except Exception as e:
        return handle_node_error(state, e, "create_document_files")

async def consolidate_final_deliverables(state: MarketingDocumentState) -> MarketingDocumentState:
    """Consolidate all results and create final deliverables package"""
    
    try:
        logger.info("Consolidating final deliverables")
        
        state = ensure_state_initialization(state)
        
        state["status"] = "consolidating"
        state["current_stage"] = "final_consolidation"
        state["updated_at"] = datetime.utcnow()
        
        # Calculate overall metrics
        total_docs = len(state["document_plans"])
        completed_docs = len([s for s in state["document_status"].values() if s == DocumentStatus.COMPLETED])
        avg_quality = sum(state["quality_scores"].values()) / len(state["quality_scores"]) if state["quality_scores"] else 0
        
        # Create executive summary
        executive_summary = {
            "project_status": "completed" if completed_docs == total_docs else "partially_completed",
            "documents_delivered": f"{completed_docs}/{total_docs}",
            "average_quality_score": round(avg_quality, 2),
            "validation_status": state.get("overall_validation_status", ValidationStatus.NOT_VALIDATED).value,
            "success_criteria_met": sum(state["criteria_validation"].values()),
            "total_success_criteria": len(state["success_criteria"]),
            "key_deliverables": [
                {
                    "title": item["title"],
                    "format": item["format"],
                    "status": item["validation_status"]
                }
                for item in state["document_manifest"][:5]  # Top 5 deliverables
            ]
        }
        
        # Create final deliverables package
        final_deliverables = {
            "request_id": state.get("request_id"),
            "correlation_id": state.get("correlation_id"),
            "completion_timestamp": datetime.utcnow().isoformat(),
            "executive_summary": executive_summary,
            "document_manifest": state["document_manifest"],
            "document_urls": state["document_urls"],
            "validation_summary": {
                "individual_validation": {
                    doc_id: {
                        "quality_score": state["quality_scores"].get(doc_id, 0),
                        "issues_found": len(state["validation_results"].get(doc_id, {}).get("issues_found", []))
                    }
                    for doc_id in state["document_contents"].keys()
                },
                "cross_document_validation": state["cross_document_validation"]
            },
            "success_criteria_validation": state["criteria_validation"],
            "processing_metadata": {
                "total_processing_time": (datetime.utcnow() - state.get("created_at", datetime.utcnow())).total_seconds(),
                "documents_generated": completed_docs,
                "validation_performed": state.get("overall_validation_status") != ValidationStatus.NOT_VALIDATED
            }
        }
        
        # Update state
        state["final_deliverables"] = final_deliverables
        state["executive_summary"] = executive_summary
        
        # Calculate processing duration
        if "created_at" in state:
            state["processing_duration"] = (datetime.utcnow() - state["created_at"]).total_seconds()
        
        state["status"] = "completed"
        state["current_stage"] = "delivery"
        state["updated_at"] = datetime.utcnow()
        
        state["messages"].append(AIMessage(
            content=f"Project complete! Delivered {completed_docs} documents with average quality score of {avg_quality:.2f}."
        ))
        
        logger.info(f"Final deliverables consolidated: {completed_docs}/{total_docs} documents delivered")
        
        return state
        
    except Exception as e:
        return handle_node_error(state, e, "consolidate_final_deliverables")

# Build the Enhanced Marketing Document Workflow
def build_marketing_workflow() -> StateGraph:
    """Build the complete marketing document generation workflow"""
    
    workflow = StateGraph(MarketingDocumentState)
    
    # Add all workflow nodes
    workflow.add_node("analyze_requirements", analyze_document_requirements)
    workflow.add_node("generate_documents", generate_all_documents)
    workflow.add_node("validate_individual", validate_individual_documents)
    workflow.add_node("validate_consistency", validate_cross_document_consistency)
    workflow.add_node("create_files", create_document_files)
    workflow.add_node("consolidate", consolidate_final_deliverables)
    
    # Set entry point
    workflow.set_entry_point("analyze_requirements")
    
    # Add workflow edges
    workflow.add_edge("analyze_requirements", "generate_documents")
    workflow.add_edge("generate_documents", "validate_individual")
    workflow.add_edge("validate_individual", "validate_consistency")
    workflow.add_edge("validate_consistency", "create_files")
    workflow.add_edge("create_files", "consolidate")
    workflow.add_edge("consolidate", END)
    
    return workflow

# Export the enhanced workflow
marketing_document_graph = build_marketing_workflow()

# Health check function
async def health_check() -> Dict[str, Any]:
    """Health check endpoint for monitoring"""
    try:
        return {
            "status": "healthy",
            "timestamp": datetime.utcnow().isoformat(),
            "version": "3.0.0",
            "capabilities": {
                "document_generation": True,
                "quality_validation": True,
                "cross_document_validation": True,
                "file_formats": ["docx", "pdf", "json"],
                "parallel_processing": True,
                "dependency_resolution": True
            },
            "workflow_stages": [
                "analyze_requirements",
                "generate_documents",
                "validate_individual",
                "validate_consistency",
                "create_files",
                "consolidate"
            ]
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat()
        }
